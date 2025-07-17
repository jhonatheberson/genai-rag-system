import io
from typing import List, Dict, Tuple
import utils
from PyPDF2 import PdfReader
from bs4 import BeautifulSoup
from docx import Document

def process_document(uploaded_file) -> Tuple[List[str], Dict]:
    """Process uploaded document and return chunks with metadata."""
    try:
        content = ""
        metadata = {
            'filename': uploaded_file.name,
            'file_size': uploaded_file.size,
            'upload_time': None  # Will be set by db_service
        }
        
        if uploaded_file.type == "application/pdf":
            content, pdf_meta = process_pdf(uploaded_file)
            metadata.update(pdf_meta)
        elif uploaded_file.type == "text/plain":
            content = uploaded_file.getvalue().decode("utf-8")
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            content, docx_meta = process_docx(uploaded_file)
            metadata.update(docx_meta)
        elif uploaded_file.type == "text/html":
            content = process_html(uploaded_file)
        elif uploaded_file.type in ["text/csv", "application/csv"]:
            content, csv_meta = process_csv(uploaded_file)
            metadata.update(csv_meta)
        else:
            raise ValueError("Unsupported file type")
        
        # Clean and chunk the content
        content = utils.sanitize_text(content)
        chunks = split_into_chunks(content)
        
        return chunks, metadata
        
    except Exception as e:
        raise ValueError(f"Error processing document: {str(e)}")

def process_pdf(file) -> Tuple[str, Dict]:
    """Extract text and metadata from PDF file."""
    try:
        pdf = PdfReader(io.BytesIO(file.getvalue()))
        content = ""
        for page in pdf.pages:
            content += page.extract_text() + "\n"
            
        metadata = {
            'format': 'pdf',
            'pages': len(pdf.pages),
            'size': len(content)
        }
        
        return content, metadata
        
    except Exception as e:
        raise ValueError(f"Error processing PDF: {str(e)}")

def process_docx(file) -> Tuple[str, Dict]:
    """Extract text and metadata from DOCX file."""
    try:
        doc = Document(io.BytesIO(file.getvalue()))
        content = ""
        for para in doc.paragraphs:
            content += para.text + "\n"
            
        metadata = {
            'format': 'docx',
            'paragraphs': len(doc.paragraphs),
            'size': len(content)
        }
        
        return content, metadata
        
    except Exception as e:
        raise ValueError(f"Error processing DOCX: {str(e)}")

def process_html(file) -> str:
    """Extract text from HTML file."""
    html_content = file.getvalue().decode("utf-8")
    soup = BeautifulSoup(html_content, 'html.parser')
    # Remove script and style elements
    for script in soup(["script", "style"]):
        script.decompose()
    text = soup.get_text(separator=" ")
    return text

def process_csv(file) -> Tuple[str, Dict]:
    """Extract text and metadata from CSV file."""
    try:
        # Read the CSV as text
        content = file.getvalue().decode("utf-8")
        
        # Create metadata
        metadata = {
            'format': 'csv',
            'rows': content.count('\n') + 1,
            'size': len(content)
        }
        
        return content, metadata
        
    except Exception as e:
        raise ValueError(f"Não foi possível processar o arquivo CSV: {str(e)}")

def split_into_chunks(text: str) -> List[str]:
    """Split text into chunks of appropriate size."""
    chunks = []
    current_chunk = ""
    
    for sentence in utils.split_into_sentences(text):
        if len(current_chunk) + len(sentence) <= utils.MAX_CHUNK_SIZE:
            current_chunk += sentence + " "
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = sentence + " "
    
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return chunks